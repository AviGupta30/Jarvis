"""
Jarvis Tool Registry — All callable actions Jarvis can perform.
Includes agentic tools for multi-step automation (PDF, Word, window mgmt, Copilot).
"""
from datetime import datetime
import webbrowser
import threading
import urllib.parse
import tkinter as tk
import subprocess
import os
import sys
import re
import pyperclip
import pyautogui
import psutil

from app.services.whatsapp import open_whatsapp, send_whatsapp_message

pyautogui.FAILSAFE = False

# ── Neural Cache client (optional — graceful degradation if server is down) ────
try:
    from neural_cache.client import CacheClient as _CacheClient
    _nc = _CacheClient()   # module-level singleton; auto-reconnects
except Exception:
    _nc = None

# ── Screen Reading (via screen_reader.py + VLM) ───────────────────────────

def read_my_screen(user_query: str = "What am I looking at?", intent_mode: str = None) -> str:
    """
    Read and describe the current screen using the VLM pipeline.
    Routes to the correct mode:
      intent_mode='describe' → structured scene summary
      intent_mode='suggest'  → 3 next best actions
      intent_mode='execute'  → step-by-step action plan
    Falls back to OCR/accessibility tree if VLM is unavailable.
    """
    try:
        from app.services.screen_reader import read_screen_as_tool
        return read_screen_as_tool(user_query=user_query)
    except Exception as e:
        return f"Could not read screen: {e}"

def play_video_in_browser() -> str:
    """
    Play a video currently visible in the active browser tab.
    Used when user says 'play that video', 'play it', 'play the 52 minute one', etc.
    Tries multiple strategies in order:
      1. Focus browser window
      2. Press Space (plays video if player is focused)
      3. Press 'k' (YouTube keyboard shortcut for play/pause)
      4. Tab to first video result and press Enter
    """
    import time
    try:
        import pygetwindow as gw
        # Find and focus browser window
        browser_names = ['chrome', 'edge', 'firefox', 'browser', 'youtube']
        browser_win = None
        for win in gw.getAllWindows():
            if win.title and any(b in win.title.lower() for b in browser_names):
                browser_win = win
                break

        if browser_win:
            if browser_win.isMinimized:
                browser_win.restore()
            browser_win.activate()
            time.sleep(0.5)
    except Exception:
        pass

    time.sleep(0.3)

    # Strategy 1: If YouTube is open and video player is visible, 'k' toggles play/pause
    # Strategy 2: Space plays/pauses most video players
    # Try clicking the center of screen first (often lands on video player)
    try:
        screen_w, screen_h = pyautogui.size()
        # Click upper-center — where the first YouTube search result typically is
        pyautogui.click(screen_w // 2, int(screen_h * 0.30))
        time.sleep(0.3)
        # Press 'k' — YouTube play/pause shortcut
        pyautogui.press('k')
        time.sleep(0.5)
        return "Pressed play on the video."
    except Exception:
        pass

    # Strategy 3: Tab to first video thumbnail and press Enter
    try:
        # Press Tab several times to reach the first video result
        for _ in range(3):
            pyautogui.press('tab')
            time.sleep(0.15)
        pyautogui.press('enter')
        return "Navigated to and opened the video."
    except Exception as e:
        return f"Could not play video: {e}"

# ─── Web & Information ──────────────────────────────────────────────────────

def _extract_location(query: str) -> str:
    """Pulls a clean location name out of a weather query."""
    # Remove filler words
    for word in ['current', 'weather', 'temperature', 'temp', 'forecast',
                 'tell me', 'what is', 'what\'s', 'the', 'whats', 'today',
                 'in', 'of', 'at', 'right now', 'now', 'india']:
        query = re.sub(rf'\b{re.escape(word)}\b', ' ', query, flags=re.IGNORECASE)
    return query.strip(' ,.-') or "Delhi"

def get_weather(location: str) -> str:
    """Get current weather using wttr.in — returns a clean spoken string."""
    try:
        import requests
        # Detailed JSON weather
        url = f"https://wttr.in/{urllib.parse.quote(location)}?format=j1"
        resp = requests.get(url, timeout=6)
        if resp.status_code == 200:
            data = resp.json()
            current = data['current_condition'][0]
            temp_c = current['temp_C']
            feels_c = current['FeelsLikeC']
            desc = current['weatherDesc'][0]['value']
            humidity = current['humidity']
            area = data['nearest_area'][0]['areaName'][0]['value']
            region = data['nearest_area'][0]['region'][0]['value']
            return (f"{desc}, {temp_c}°C (feels like {feels_c}°C), "
                    f"humidity {humidity}% in {area}, {region}.")
    except Exception:
        pass
    # Fallback: simple format string
    try:
        import requests
        resp = requests.get(f"https://wttr.in/{urllib.parse.quote(location)}?format=3", timeout=5)
        if resp.status_code == 200:
            return resp.text.strip()
    except Exception:
        pass
    return f"Couldn't fetch weather for {location}."

def get_info(query: str) -> str:
    """
    Smart information lookup (Step 3 upgrade):
    - Weather queries  -> wttr.in (real-time, reliable)
    - Everything else  -> web_search.smart_search() (DuckDuckGo + Wikipedia fallback)
    Does NOT open any browser window.
    """
    lower = query.lower()

    # Route weather queries to the dedicated weather API
    weather_kw = ['weather', 'temperature', 'temp', 'rain', 'forecast',
                  'humidity', 'mausam', 'barish']
    if any(kw in lower for kw in weather_kw):
        location = _extract_location(query)
        return get_weather(location)

    # All other queries -> web_search module (DuckDuckGo primary, Wikipedia fallback)
    try:
        from app.services.web_search import smart_search
        return smart_search(query)
    except Exception as e:
        return f"Search failed: {e}"

SITE_SHORTCUTS = {
    # Global
    "google": "https://www.google.com",
    "youtube": "https://www.youtube.com",
    "github": "https://www.github.com",
    "gmail": "https://mail.google.com",
    "maps": "https://maps.google.com",
    "reddit": "https://www.reddit.com",
    "wikipedia": "https://www.wikipedia.org",
    "twitter": "https://www.twitter.com",
    "x": "https://www.x.com",
    "instagram": "https://www.instagram.com",
    "linkedin": "https://www.linkedin.com",
    "netflix": "https://www.netflix.com",
    "spotify": "https://open.spotify.com",
    "amazon": "https://www.amazon.in",
    "whatsapp": "whatsapp:",
    "chatgpt": "https://chat.openai.com",
    "gemini": "https://gemini.google.com",
    # Indian streaming
    "hotstar": "https://www.hotstar.com",
    "disney": "https://www.hotstar.com",
    "disney hotstar": "https://www.hotstar.com",
    "jiocinema": "https://www.jiocinema.com",
    "jio cinema": "https://www.jiocinema.com",
    "jio": "https://www.jiocinema.com",
    "zee5": "https://www.zee5.com",
    "sonyliv": "https://www.sonyliv.com",
    "sony liv": "https://www.sonyliv.com",
    "prime": "https://www.primevideo.com",
    "prime video": "https://www.primevideo.com",
    "amazon prime": "https://www.primevideo.com",
    "mxplayer": "https://www.mxplayer.in",
    # Social & productivity
    "facebook": "https://www.facebook.com",
    "telegram": "https://web.telegram.org",
    "notion": "https://www.notion.so",
    "figma": "https://www.figma.com",
    "canva": "https://www.canva.com",
    "stackoverflow": "https://www.stackoverflow.com",
    "stack overflow": "https://www.stackoverflow.com",
    "flipkart": "https://www.flipkart.com",
    "swiggy": "https://www.swiggy.com",
    "zomato": "https://www.zomato.com",
}

def open_safe_website(url: str, browser: str = None) -> str:
    """Opens a specific website in the browser."""
    url_lower = url.lower().strip()
    
    # Strip accidental www. prefix for shortcut matching
    clean_url = re.sub(r'^www\.', '', url_lower)
    
    if clean_url in SITE_SHORTCUTS:
        final_url = SITE_SHORTCUTS[clean_url]
    elif url.startswith("http"):
        final_url = url
    elif "." in clean_url and not any(space in clean_url for space in [' ', '\n', '\t']):
        final_url = f"https://{clean_url}"
    else:
        final_url = f"https://www.{clean_url.replace(' ', '')}.com"
        
    if browser and "chrome" in browser.lower():
        subprocess.Popen(f'start chrome "{final_url}"', shell=True)
        return f"Opened {final_url} in Chrome"
    elif browser and "edge" in browser.lower():
        subprocess.Popen(f'start msedge "{final_url}"', shell=True)
        return f"Opened {final_url} in Edge"
        
    webbrowser.open(final_url)
    return f"Opened: {final_url}"

def open_google_search_in_browser(query: str) -> str:
    """Opens Google search in browser for user to see results."""
    encoded = urllib.parse.quote_plus(query)
    webbrowser.open(f"https://www.google.com/search?q={encoded}")
    return f"Searched Google for: {query}"

def youtube_search(query: str, autoplay: bool = False) -> str:
    """Search YouTube and either play the first video or show results."""
    encoded = urllib.parse.quote_plus(query)
    if autoplay:
        try:
            from urllib.request import urlopen
            html = urlopen(f"https://www.youtube.com/results?search_query={encoded}").read().decode()
            video_ids = re.findall(r"watch\?v=([a-zA-Z0-9_-]{11})", html)
            if video_ids:
                first_vid = video_ids[0]
                webbrowser.open(f"https://www.youtube.com/watch?v={first_vid}")
                return f"Playing first YouTube result for: {query}"
        except Exception as e:
            pass
    
    # Fallback or if autoplay is False
    webbrowser.open(f"https://www.youtube.com/results?search_query={encoded}")
    return f"Searching YouTube for: {query}"

# ─── Date & Time ────────────────────────────────────────────────────────────

def get_system_time() -> str:
    """Returns current date and time naturally."""
    return datetime.now().strftime("%I:%M %p on %A, %B %d, %Y")

# ─── System Information ─────────────────────────────────────────────────────

def get_system_info() -> str:
    """Returns CPU usage, RAM usage, and battery status."""
    cpu = psutil.cpu_percent(interval=1)
    ram = psutil.virtual_memory()
    ram_used = ram.used // (1024**3)
    ram_total = ram.total // (1024**3)
    info = f"CPU: {cpu}% | RAM: {ram_used}GB of {ram_total}GB used"
    try:
        battery = psutil.sensors_battery()
        if battery:
            status = "charging" if battery.power_plugged else "on battery"
            info += f" | Battery: {int(battery.percent)}% ({status})"
    except Exception:
        pass
    return info

# ─── Screen & Window Control ────────────────────────────────────────────────

def take_screenshot(filename: str = None) -> str:
    """Takes a full screenshot and saves to Desktop."""
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    fname = filename or f"screenshot_{datetime.now().strftime('%H%M%S')}.png"
    save_path = os.path.join(desktop, fname)
    pyautogui.screenshot().save(save_path)
    return "Screenshot saved to Desktop."

def snap_windows(left_app: str, right_app: str) -> str:
    """Snaps left_app to left half and right_app to right half using Win32 API (no keyboard shortcuts)."""
    from app.services.window_layout import win32_snap_two_windows
    return win32_snap_two_windows(left_app, right_app)

# ─── File Operations ────────────────────────────────────────────────────────
def create_folder(folder_name: str) -> str:
    """Creates a new folder on the Desktop."""
    try:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        path = os.path.join(desktop, folder_name)
        os.makedirs(path, exist_ok=True)
        return f"Successfully created folder '{folder_name}' on Desktop."
    except Exception as e:
        return f"Failed to create folder: {str(e)}"

def create_file(filepath: str, content: str) -> str:
    """Creates a new file with the given content."""
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return f"Successfully created file: {filepath}"
    except Exception as e:
        return f"Failed to create file: {str(e)}"

def append_to_file(filepath: str, content: str) -> str:
    """Appends content to an existing file."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f:
            f.write("\n" + content)
        return f"Successfully appended to file: {filepath}"
    except Exception as e:
        return f"Failed to edit file: {str(e)}"

def close_tab() -> str:
    """Closes the current browser tab using Ctrl+W, targeting the real foreground app."""
    from app.services.window_layout import win32_close_active_tab
    return win32_close_active_tab()

def close_window() -> str:
    """Closes the current active real app window (skips the Jarvis overlay)."""
    from app.services.window_layout import win32_close_active_window
    return win32_close_active_window()

def _find_window_fuzzy(app_name: str):
    """Find a window HWND by fuzzy name matching using Win32 API (no pygetwindow)."""
    from app.services.window_layout import win32_find_window
    return win32_find_window(app_name)  # returns int HWND or None

def close_specific_window(app_name: str) -> str:
    """Closes a specific window/app by name using Win32 PostMessage WM_CLOSE."""
    from app.services.window_layout import win32_close_window
    return win32_close_window(app_name)

def minimize_window(app_name: str) -> str:
    """Minimizes a specific window by name using Win32 ShowWindow."""
    from app.services.window_layout import win32_minimize_window
    return win32_minimize_window(app_name)

def maximize_window(app_name: str) -> str:
    """Maximizes a specific window by name using Win32 ShowWindow."""
    from app.services.window_layout import win32_maximize_window
    return win32_maximize_window(app_name)

def minimize_all_windows() -> str:
    """Minimizes all open windows to show the desktop."""
    pyautogui.hotkey('win', 'd')
    return "All windows minimized."

def lock_screen() -> str:
    """Locks the Windows screen."""
    pyautogui.hotkey('win', 'l')
    return "Screen locked."

# ─── Volume & Media Control ─────────────────────────────────────────────────

def volume_up(steps: int = 5) -> str:
    """Increases system volume."""
    for _ in range(steps):
        pyautogui.press('volumeup')
    return f"Volume increased."

def volume_down(steps: int = 5) -> str:
    """Decreases system volume."""
    for _ in range(steps):
        pyautogui.press('volumedown')
    return f"Volume decreased."

def mute_volume() -> str:
    """Toggles mute on the system."""
    pyautogui.press('volumemute')
    return "Volume toggled mute."

def media_play_pause() -> str:
    """Plays or pauses media (Spotify, YouTube, etc.)."""
    pyautogui.press('playpause')
    return "Media play/pause toggled."

def media_next() -> str:
    """Skips to next track."""
    pyautogui.press('nexttrack')
    return "Skipped to next track."

def media_previous() -> str:
    """Goes to previous track."""
    pyautogui.press('prevtrack')
    return "Went to previous track."

def play_music(song: str) -> str:
    """Play a song on Spotify by searching and playing the top result."""
    from app.services.spotify_service import play_song_dynamic
    import threading

    def _do_play():
        try:
            play_song_dynamic(song)
        except Exception as e:
            print(f"[Jarvis] Spotify automation failed: {e}")

    threading.Thread(target=_do_play, daemon=True).start()
    return f"Playing '{song}' on Spotify."

# ─── Clipboard & Typing ─────────────────────────────────────────────────────

def read_clipboard() -> str:
    """Reads and returns the current clipboard content."""
    content = pyperclip.paste()
    return f"Clipboard contains: {content[:500]}" if content else "Clipboard is empty."

def write_clipboard(text: str) -> str:
    """Copies text to the clipboard."""
    pyperclip.copy(text)
    return f"Copied to clipboard: '{text}'"

def type_text(text: str) -> str:
    """Types the given text into the currently focused window."""
    import time
    time.sleep(0.5)
    pyperclip.copy(text)
    pyautogui.hotkey('ctrl', 'v')
    return f"Typed: '{text}'"

# ─── App Launcher ────────────────────────────────────────────────────────────

def open_app(app_name: str) -> str:
    """Opens a Windows application by name."""
    apps = {
        "notepad": "notepad.exe", "calculator": "calc.exe", "paint": "mspaint.exe",
        "file explorer": "explorer.exe", "explorer": "explorer.exe",
        "task manager": "taskmgr.exe", "cmd": "cmd.exe", "command prompt": "cmd.exe",
        "terminal": "wt.exe", "vs code": "code", "vscode": "code",
        "word": "winword.exe", "excel": "excel.exe", "powerpoint": "powerpnt.exe",
        "chrome": "chrome.exe", "edge": "msedge.exe",
        "spotify": "spotify.exe", "discord": "discord.exe", "zoom": "zoom.exe",
        "settings": "ms-settings:", "control panel": "control.exe",
        "snipping tool": "snippingtool.exe",
    }
    exe = apps.get(app_name.lower().strip(), app_name)
    try:
        subprocess.Popen(exe, shell=True)
        return f"Opened {app_name}."
    except Exception as e:
        return f"Could not open {app_name}: {e}"

# ─── Sticky Notes ────────────────────────────────────────────────────────────

def _spawn_sticky(content: str):
    root = tk.Tk()
    root.title("📌 Jarvis Note")
    screen_w = root.winfo_screenwidth()
    root.geometry(f"300x200+{screen_w - 320}+40")
    root.configure(bg="#FFEF9F")
    root.attributes("-topmost", True)
    header = tk.Frame(root, bg="#F5C518", height=28)
    header.pack(fill=tk.X)
    tk.Label(header, text="📌 Jarvis Note", bg="#F5C518",
             font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT, padx=8, pady=4)
    tk.Button(header, text="✕", bg="#F5C518", relief=tk.FLAT,
              command=root.destroy).pack(side=tk.RIGHT, padx=6, pady=2)
    tk.Text(root, bg="#FFEF9F", fg="#1a1a1a", font=("Segoe UI", 11),
            wrap=tk.WORD, relief=tk.FLAT, bd=8).pack(fill=tk.BOTH, expand=True)
    root.nametowidget(root.winfo_children()[-1]).insert(tk.END, content)
    root.mainloop()

def create_sticky_note(content: str) -> str:
    """Creates a floating sticky note on screen."""
    threading.Thread(target=_spawn_sticky, args=(content,), daemon=True).start()
    return f"Sticky note created."

def close_sticky_notes() -> str:
    """Closes all active sticky notes created by Jarvis."""
    try:
        import pygetwindow as gw
        closed = 0
        for w in gw.getAllWindows():
            if w.title and "📌 Jarvis Note" in w.title:
                w.close()
                closed += 1
        if closed > 0:
            return f"Closed {closed} sticky note(s)."
        return "No sticky notes found."
    except Exception as e:
        return f"Failed to close sticky notes: {e}"

# ─── Math & Calculations ─────────────────────────────────────────────────────

def calculate(expression: str) -> str:
    """Safely evaluates a math expression and returns the result."""
    try:
        # Only allow safe math characters
        safe = re.sub(r'[^0-9+\-*/.() %]', '', expression)
        result = eval(safe)
        return f"{expression} = {result}"
    except Exception as e:
        return f"Could not calculate: {expression} ({e})"

# ─── Reminders ───────────────────────────────────────────────────────────────

def set_reminder(message: str, seconds: int = 60) -> str:
    """Sets a reminder that Jarvis will speak after a given number of seconds."""
    def _remind():
        import time
        time.sleep(seconds)
        # Import here to avoid circular import
        import asyncio
        from app.services.voice import speak_text as _speak
        asyncio.run(_speak(f"Reminder: {message}"))

    threading.Thread(target=_remind, daemon=True).start()
    minutes = seconds // 60
    return f"Reminder set for {minutes} minute{'s' if minutes != 1 else ''}: '{message}'"

# ─── Memory Tools ─────────────────────────────────────────────────────────────

def remember_preference(key: str, value: str) -> str:
    """Save a user preference to long-term memory."""
    try:
        from app.memory import save_preference
        save_preference(key.strip(), value.strip())
        return f"Got it! I'll remember that your {key} preference is: {value}."
    except Exception as e:
        return f"Couldn't save preference: {e}"

def list_learned_skills() -> str:
    """Return all dynamic skills Jarvis has learned."""
    try:
        from app.memory import list_skills
        skills = list_skills()
        if not skills:
            return "I haven't learned any custom skills yet. Give me tasks to figure out and I'll remember them!"
        lines = "\n".join(f"• {s}" for s in skills[:20])
        return f"I've learned {len(skills)} custom skill(s):\n{lines}"
    except Exception as e:
        return f"Couldn't retrieve skills: {e}"

# ─── Agentic / Multi-Step Tools ───────────────────────────────────────────────

def read_pdf_text(path: str = None, filename: str = None) -> str:
    """Extract all text from a PDF file. Accepts full path or filename to search on Desktop."""
    import time
    # Resolve path if only filename given
    if not path and filename:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        docs = os.path.join(os.path.expanduser("~"), "Documents")
        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
        for folder in [desktop, docs, downloads]:
            candidate = os.path.join(folder, filename)
            if os.path.exists(candidate):
                path = candidate
                break
        if not path:
            # Try case-insensitive search
            for folder in [desktop, docs, downloads]:
                try:
                    for f in os.listdir(folder):
                        if f.lower() == filename.lower():
                            path = os.path.join(folder, f)
                            break
                except Exception:
                    pass
                if path:
                    break
    if not path or not os.path.exists(path):
        return f"PDF file not found: {filename or path}"
    # Try pdfplumber first (best for text PDFs)
    try:
        import pdfplumber
        text_pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_pages.append(t)
        text = "\n\n".join(text_pages)
        if text.strip():
            return text[:8000] # cap at 8000 chars
    except ImportError:
        pass
    except Exception as e:
        pass
    # Try PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        text = "\n\n".join(page.get_text() for page in doc)
        if text.strip():
            return text[:8000]
    except ImportError:
        pass
    except Exception:
        pass
    # Fallback: install pdfplumber and retry
    try:
        subprocess.run([sys.executable if 'sys' in dir() else 'python', '-m', 'pip', 'install', 'pdfplumber', '-q'], timeout=30)
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n\n".join(p.extract_text() or "" for p in pdf.pages)
        return text[:8000] if text.strip() else "Could not extract text from PDF."
    except Exception as e:
        return f"Failed to read PDF: {e}"

def find_file(name: str, location: str = "all") -> str:
    """Find a file by name (or partial name) on Desktop, Documents, Downloads."""
    search_dirs = {
        "desktop": [os.path.join(os.path.expanduser("~"), "Desktop")],
        "documents": [os.path.join(os.path.expanduser("~"), "Documents")],
        "downloads": [os.path.join(os.path.expanduser("~"), "Downloads")],
        "all": [
            os.path.join(os.path.expanduser("~"), "Desktop"),
            os.path.join(os.path.expanduser("~"), "Documents"),
            os.path.join(os.path.expanduser("~"), "Downloads"),
        ]
    }
    dirs = search_dirs.get(location.lower(), search_dirs["all"])
    name_lower = name.lower()
    found = []
    for d in dirs:
        try:
            for f in os.listdir(d):
                if name_lower in f.lower():
                    found.append(os.path.join(d, f))
        except Exception:
            pass
    if not found:
        return f"No file matching '{name}' found in {location}."
    return "\n".join(found[:5])

def open_file(path: str) -> str:
    """Open any file with its default application."""
    try:
        os.startfile(path)
        return f"Opened: {os.path.basename(path)}"
    except Exception as e:
        try:
            subprocess.Popen(['cmd', '/c', 'start', '', path], shell=False)
            return f"Opened: {os.path.basename(path)}"
        except Exception as e2:
            return f"Failed to open file: {e2}"

def focus_window(name: str, timeout: float = 5.0) -> str:
    """Bring a window to the foreground by partial name match (Win32 API, no pygetwindow)."""
    import time
    from app.services.window_layout import win32_focus_window, win32_find_window
    deadline = time.time() + timeout
    while time.time() < deadline:
        result = win32_focus_window(name)
        if result.startswith('✅'):
            return result
        time.sleep(0.4)
    return f"Window '{name}' not found after {timeout}s."

def wait_for_window(name: str, timeout: float = 15.0) -> str:
    """Block until a window with the given name appears, then focus it."""
    import time
    from app.services.window_layout import win32_focus_window
    deadline = time.time() + timeout
    while time.time() < deadline:
        result = win32_focus_window(name)
        if result.startswith('✅'):
            return f"Window appeared and focused: {result}"
        time.sleep(0.5)
    return f"Timed out waiting for window '{name}' after {timeout}s."

def type_and_submit(text: str, delay: float = 0.3) -> str:
    """Type text into the focused field, then press Enter."""
    import time
    time.sleep(delay)
    pyperclip.copy(text)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(0.3)
    pyautogui.press('enter')
    return f"Typed and submitted: '{text[:60]}'"

def copy_selected_text() -> str:
    """Select all text in the focused element and copy to clipboard, then return it."""
    import time
    pyautogui.hotkey('ctrl', 'a')
    time.sleep(0.3)
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(0.5)
    content = pyperclip.paste()
    return content[:5000] if content else "(clipboard empty after copy)"

def create_word_doc(filename: str, content: str, save_path: str = None) -> str:
    """Create a .docx Word document with the given content.
    content can be a list of (heading, body) tuples formatted as 'Q: ...\nA: ...' blocks.
    """
    if not save_path:
        save_path = os.path.join(os.path.expanduser("~"), "Desktop", filename)
    if not save_path.endswith('.docx'):
        save_path += '.docx'
    try:
        try:
            from docx import Document
        except ImportError:
            subprocess.run([sys.executable if 'sys' in dir() else 'python', '-m', 'pip', 'install', 'python-docx', '-q'], timeout=60)
            from docx import Document
        doc = Document()
        doc.add_heading(os.path.splitext(os.path.basename(save_path))[0], 0)
        # Parse Q&A blocks if present
        blocks = content.split('\n\n')
        for block in blocks:
            lines = block.strip().split('\n')
            for line in lines:
                if line.startswith('Q:') or line.startswith('Question:'):
                    doc.add_heading(line, level=2)
                elif line.startswith('A:') or line.startswith('Answer:'):
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph(line)
        doc.save(save_path)
        return f"Word document saved: {save_path}"
    except Exception as e:
        # Fallback: save as plain text .txt
        txt_path = save_path.replace('.docx', '.txt')
        try:
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return f"Saved as text file (python-docx unavailable): {txt_path}"
        except Exception as e2:
            return f"Failed to create document: {e} | txt fallback: {e2}"

def read_active_window_text() -> str:
    """Get all readable text from the currently active window."""
    try:
        from app.services.ui_inspector import get_active_window_info
        return get_active_window_info()
    except Exception as e:
        return f"Could not read active window: {e}"

# ── UIA-powered automation tools ───────────────────────────────────────────────────

def click_ui_element_uia(app_title: str, element_name: str = None,
                         automation_id: str = None,
                         control_type: str = None) -> str:
    """
    Click a UI element inside an app by AutomationId, name, or control type.
    Does NOT move the physical mouse cursor — fires via Windows UI Automation.
    Preferred over coordinate-based clicking for any known app element.

    Args:
        app_title:     Partial window title of the target app (e.g. 'VS Code')
        element_name:  Visible label of the element (e.g. 'Run Test')
        automation_id: UIA AutomationId — most stable identifier (e.g. 'workbench.panel.terminal')
        control_type:  UIA type: 'Button', 'Edit', 'MenuItem', 'Pane', 'CheckBox', etc.
    """
    try:
        from app.services.ui_inspector import smart_click
        return smart_click(app_title, element_name=element_name,
                           automation_id=automation_id, control_type=control_type)
    except Exception as e:
        return f"UIA click failed: {e}"

def type_into_ui_element(app_title: str, element_name: str = None,
                         text: str = "", automation_id: str = None) -> str:
    """
    Inject text into a specific input field in an app via UIA Value pattern.
    No simulated global keystrokes — sets the value directly in the OS control.
    Works even when the window is behind other windows.

    Args:
        app_title:     Partial window title
        element_name:  Label of the text field (e.g. 'Terminal input')
        text:          Text to enter
        automation_id: Optional AutomationId for maximum stability
    """
    try:
        from app.services.ui_inspector import type_into_element
        return type_into_element(app_title, element_name=element_name,
                                 text=text, automation_id=automation_id)
    except Exception as e:
        return f"UIA type failed: {e}"

def read_ui_element_text(app_title: str, element_name: str = None,
                          automation_id: str = None) -> str:
    """
    Read the current text content of a UI element — e.g. a terminal output
    pane, a label, or a text field. Uses UIA TextPattern.

    Args:
        app_title:     Partial window title
        element_name:  Name of the element to read
        automation_id: Optional AutomationId
    """
    try:
        from app.services.ui_inspector import read_element_text
        return read_element_text(app_title, element_name=element_name,
                                 automation_id=automation_id)
    except Exception as e:
        return f"UIA read failed: {e}"

def dump_app_ui_tree(app_title: str, depth: int = 3) -> str:
    """
    Dump the full Windows UI Automation accessibility tree of an app window.
    Use this once per app to discover AutomationIds for stable element targeting.
    Output is a human-readable indented tree showing all buttons, inputs, panes.

    Args:
        app_title: Partial window title (e.g. 'Visual Studio Code', 'Discord')
        depth:     How many levels deep to walk (default 3, max 5)
    """
    try:
        from app.services.ui_inspector import debug_ui_tree
        return debug_ui_tree(app_title, depth=depth)
    except Exception as e:
        return f"UIA tree dump failed: {e}"


def open_windows_copilot() -> str:
    """Open Windows Copilot sidebar using Win+C hotkey."""
    import time
    pyautogui.hotkey('win', 'c')
    time.sleep(2.5)  # Wait for Copilot panel to animate open
    return "Windows Copilot opened."

def send_to_copilot(question: str, wait_seconds: float = 8.0) -> str:
    """Type a question into the Windows Copilot sidebar and retrieve the response.
    Assumes Copilot is already open and focused.
    """
    import time
    # Click Copilot input area — it's usually at the bottom of the panel
    # Use the UI inspector to find the text box
    try:
        from app.services.ui_inspector import click_ui_element
        clicked = click_ui_element("Ask me anything")
        if not clicked:
            clicked = click_ui_element("Message Copilot")
        if not clicked:
            # Try clicking in the lower-right area where Copilot input usually sits
            screen_w, screen_h = pyautogui.size()
            pyautogui.click(screen_w - 200, screen_h - 80)
        time.sleep(0.5)
    except Exception:
        pass
    # Type the question
    pyperclip.copy(question)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(0.4)
    pyautogui.press('enter')
    # Wait for Copilot to respond
    time.sleep(wait_seconds)
    # Try to copy response — Copilot has a copy button after each response
    try:
        from app.services.ui_inspector import click_ui_element
        click_ui_element("Copy")
        time.sleep(0.5)
        response = pyperclip.paste()
        if response and response != question:
            return response[:3000]
    except Exception:
        pass
    # Fallback: select all text in panel and copy
    pyautogui.hotkey('ctrl', 'a')
    time.sleep(0.3)
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(0.3)
    response = pyperclip.paste()
    return response[:3000] if response else "(Could not retrieve Copilot response)"

# ── Web Search Tools (Step 3) ─────────────────────────────────────────────────

def search_site_tool(query: str, site_url: str) -> str:
    """Search for a query within a specific website using DuckDuckGo site: operator."""
    try:
        from app.services.web_search import search_site
        return search_site(query, site_url)
    except Exception as e:
        return f"Site search failed: {e}"

def scrape_url_tool(url: str) -> str:
    """Read and extract readable text content from a specific URL."""
    try:
        from app.services.web_search import scrape_url
        return scrape_url(url)
    except Exception as e:
        return f"Could not read URL: {e}"

# ── PowerPoint AI Wrappers ──────────────────────────────────────────────────
# Isolated in ppt_tool.py. These thin wrappers convert the returned dict to a
# human-readable string so it flows cleanly through the /chat response pipeline.

def _ppt_create(user_prompt: str, style: str = None, purpose: str = None):
    """
    Generator wrapper — streams live progress to the frontend via chat.py's
    inspect.isgenerator() streaming path. Each yielded string appears as a
    new line in the Jarvis chat window in real time.
    """
    try:
        from app.services.ppt_tool import ppt_create
        yield from ppt_create(prompt=user_prompt, style=style, purpose=purpose)
    except Exception as e:
        yield f"❌ PPT tool error: {e}"

def _research_and_create_ppt(topic: str, style: str = None):
    """
    Phase 3: Autonomous Research Aggregator.
    Scrapes live facts from the web, extracts them using NLP, and passes them
    to the PPT engine to create a grounded presentation with real charts.
    """
    try:
        from app.services.research_pipeline import research_topic
        from app.services.ppt_tool import ppt_create
        
        # 1. Run research pipeline (yields progress strings)
        research_data = None
        for msg in research_topic(topic):
            if isinstance(msg, str):
                yield msg
            elif isinstance(msg, dict):
                research_data = msg
        
        # 2. Run PPT creation with injected research
        yield from ppt_create(prompt=topic, style=style, research_data=research_data)
        
    except Exception as e:
        yield f"❌ Research Pipeline Error: {e}"



def _ppt_edit(edit_prompt: str):
    """
    Generator wrapper — streams live progress for slide edits.
    """
    try:
        from app.services.ppt_tool import ppt_edit
        yield from ppt_edit(edit_prompt=edit_prompt)
    except Exception as e:
        yield f"❌ PPT edit error: {e}"


def _ppt_styles() -> str:
    """List all available presentation design personalities."""
    try:
        from app.services.ppt_tool import ppt_styles
        result = ppt_styles()
        lines = [f"I have {result['count']} design styles for presentations:"]
        for k, v in result["styles"].items():
            lines.append(f"  • {v['name']} ({k}) — {v['desc']}")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ PPT styles error: {e}"


# ── WhatsApp Intelligence Tools (isolated package) ───────────────────────────

def read_whatsapp_thread(contact: str, n_messages: int = 20) -> str:
    """
    Open a WhatsApp chat, read the last N messages, and return a structured
    thread summary showing who said what and whether a reply is needed.
    Use this BEFORE generate_reply_draft to see what's going on in a chat.
    """
    try:
        from app.services.whatsapp_intelligence.thread_extractor import extract_thread_as_string
        return extract_thread_as_string(contact, n_messages)
    except Exception as e:
        return f"Could not read thread: {e}"


def build_style_profile(chat_export_path: str, your_name: str = "You") -> str:
    """
    One-time training: parse a WhatsApp .txt chat export to build your
    personal reply-style profile (length, Hinglish ratio, emoji use, etc.).
    Export any chat from WhatsApp → More → Export Chat → Without Media.
    your_name must match the name shown in the export file.
    """
    try:
        from app.services.whatsapp_intelligence.style_profiler import build_style_profile as _build
        return _build(chat_export_path, your_name)
    except Exception as e:
        return f"Could not build style profile: {e}"


def generate_reply_draft(contact: str, n_messages: int = 20) -> str:
    """
    Read the latest WhatsApp thread with a contact and generate 3 reply
    drafts that sound like YOU — matching your tone, length, Hinglish mix,
    and emoji habits from your style profile.
    After calling this, say 'send reply 1' (or 2 / 3) to send a draft.
    """
    try:
        from app.services.whatsapp_intelligence.reply_generator import generate_reply_draft as _gen
        return _gen(contact, n_messages)
    except Exception as e:
        return f"Could not generate reply drafts: {e}"


def send_style_reply(contact: str = "", draft_index: int = 1) -> str:
    """
    Send one of the reply drafts generated by generate_reply_draft.
    draft_index is 1-based: 1 = top draft, 2 = second, 3 = third.
    Also records the sent reply to improve your style profile over time.
    """
    try:
        from app.services.whatsapp_intelligence.reply_generator import send_style_reply as _send
        return _send(contact, draft_index)
    except Exception as e:
        return f"Could not send style reply: {e}"


# ── Neural Cache Tools ─────────────────────────────────────────────────────────────────────
# Isolated per Rule #1 and Rule #4. Client imported at module level above.
# Accessible via both frontend and voice through the unified /chat route.

def cache_set(key: str, value: str, ttl: int = None) -> str:
    """
    Store a key-value pair in Neural Cache (in-memory, O(1) access).
    Useful for short-lived state like session flags, counters, and mode indicators.

    Args:
        key:   Unique identifier for the value (e.g. 'focus_mode', 'current_task').
        value: The value to store (string).
        ttl:   Optional time-to-live in seconds (e.g. 1800 for 30 minutes).
               Omit or set to None for a permanent entry.

    Returns:
        Confirmation string.
    """
    try:
        if _nc is None:
            return "Neural Cache server is not running. Value not stored."
        ok = _nc.set(key, str(value), ttl=ttl)
        if ok:
            ttl_msg = f" (expires in {ttl}s)" if ttl else ""
            return f"Stored '{key}' = '{value}' in Neural Cache{ttl_msg}."
        return f"Failed to store '{key}' in Neural Cache."
    except Exception as e:
        return f"Neural Cache set error: {e}"


def cache_get(key: str) -> str:
    """
    Retrieve a value from Neural Cache.
    Returns the value if found, or a 'not found' message if the key
    doesn't exist or has expired.

    Args:
        key: The key to look up.

    Returns:
        The stored value string, or a descriptive 'not found' message.
    """
    try:
        if _nc is None:
            return "Neural Cache server is not running."
        value = _nc.get(key)
        if value is None:
            return f"No value found for '{key}' in Neural Cache (may have expired or never been set)."
        return f"Neural Cache '{key}' = '{value}'."
    except Exception as e:
        return f"Neural Cache get error: {e}"


def get_dsa_cache_status() -> str:
    """
    Check the current DSA mode status from Neural Cache.
    Returns a human-readable summary of active/inactive state and progress.
    """
    try:
        if _nc is None:
            return "Neural Cache is not available. Check DSA status directly."
        active    = _nc.get("dsa_mode:active")
        total     = _nc.get("dsa_mode:num_questions")
        completed = _nc.get("dsa_mode:completed")
        if active == "true":
            return (
                f"DSA Mode is ACTIVE. "
                f"Progress: {completed or 0}/{total or '?'} questions completed."
            )
        return "DSA Mode is currently INACTIVE."
    except Exception as e:
        return f"Could not read DSA status from cache: {e}"


TOOL_REGISTRY = {
    # Information & Web
    "get_info": get_info,

    "get_system_time": get_system_time,
    "get_system_info": get_system_info,
    "open_website": open_safe_website,
    "open_google_search_in_browser": open_google_search_in_browser,
    "youtube_search": youtube_search,
    "take_screenshot": take_screenshot,
    "snap_windows": snap_windows,
    # File Operations
    "create_folder": create_folder,
    "create_file": create_file,
    "append_to_file": append_to_file,
    # Browser controls
    "close_tab": close_tab,
    "close_window": close_window,
    "close_specific_window": close_specific_window,
    "minimize_window": minimize_window,
    "maximize_window": maximize_window,
    "minimize_all_windows": minimize_all_windows,
    "lock_screen": lock_screen,
    # Volume / Media
    "volume_up": volume_up,
    "volume_down": volume_down,
    "mute_volume": mute_volume,
    "play_music": play_music,
    "media_play_pause": media_play_pause,
    "media_next": media_next,
    "media_previous": media_previous,
    # Clipboard & Typing
    "read_clipboard": read_clipboard,
    "write_clipboard": write_clipboard,
    "type_text": type_text,
    # Apps
    "open_app": open_app,
    # Notes & Reminders
    "create_sticky_note": create_sticky_note,
    "close_sticky_notes": close_sticky_notes,
    "set_reminder": set_reminder,
    # Math
    "calculate": calculate,
    # WhatsApp (Smart — Step 11)
    "open_whatsapp":        lambda: __import__('app.services.whatsapp_smart', fromlist=['open_whatsapp']).open_whatsapp(),
    "search_whatsapp_contact": lambda name: __import__('app.services.whatsapp_smart', fromlist=['search_whatsapp_contact']).search_whatsapp_contact(name),
    "initiate_whatsapp_send":  lambda contact_name, message: __import__('app.services.whatsapp_smart', fromlist=['initiate_whatsapp_send']).initiate_whatsapp_send(contact_name, message),
    "confirm_whatsapp_send":   lambda contact_name, message: __import__('app.services.whatsapp_smart', fromlist=['confirm_whatsapp_send']).confirm_whatsapp_send(contact_name, message),
    "read_whatsapp_messages":  lambda contact_name, count=5: __import__('app.services.whatsapp_smart', fromlist=['read_whatsapp_messages']).read_whatsapp_messages(contact_name, count),
    "send_whatsapp_message":   lambda contact_name, message: __import__('app.services.whatsapp_smart', fromlist=['initiate_whatsapp_send']).initiate_whatsapp_send(contact_name, message),  # routes to confirm flow
    "initiate_whatsapp_call":  lambda contact_name: __import__('app.services.whatsapp_call', fromlist=['initiate_whatsapp_call']).initiate_whatsapp_call(contact_name),
    "confirm_whatsapp_call":   lambda contact_name: __import__('app.services.whatsapp_call', fromlist=['confirm_whatsapp_call']).confirm_whatsapp_call(contact_name),
    # WhatsApp Intelligence (Step 12 — reply style cloning)
    "read_whatsapp_thread":    read_whatsapp_thread,
    "build_style_profile":     build_style_profile,
    "generate_reply_draft":    generate_reply_draft,
    "send_style_reply":        send_style_reply,
    # Memory
    "remember_preference": remember_preference,
    "list_learned_skills": list_learned_skills,
    # ── Agentic / Multi-Step Tools ──
    "read_pdf_text": read_pdf_text,
    "find_file": find_file,
    "open_file": open_file,
    "focus_window": focus_window,
    "wait_for_window": wait_for_window,
    "adjust_active_window": lambda position=None, width_percent=None, height_percent=None, app_name=None: __import__('app.services.window_layout', fromlist=['adjust_active_window']).adjust_active_window(position, width_percent, height_percent, app_name),
    "type_and_submit": type_and_submit,
    "copy_selected_text": copy_selected_text,
    "create_word_doc": create_word_doc,
    "read_active_window_text": read_active_window_text,
    "open_windows_copilot": open_windows_copilot,
    "send_to_copilot": send_to_copilot,
    # Screen vision (Step 1)
    "read_my_screen": read_my_screen,
    # Browser video playback (bug fix)
    "play_video_in_browser": play_video_in_browser,
    # Web search (Step 3)
    "search_site": search_site_tool,
    "scrape_url": scrape_url_tool,
    # Full file system ops (Step 4)
    "read_file":       lambda path: __import__('app.services.file_ops', fromlist=['read_file']).read_file(path),
    "write_file":      lambda path, content: __import__('app.services.file_ops', fromlist=['write_file']).write_file(path, content),
    "append_file":     lambda path, content: __import__('app.services.file_ops', fromlist=['append_file']).append_file(path, content),
    "list_directory":   lambda path='Desktop': __import__('app.services.file_ops', fromlist=['list_directory']).list_directory(path),
    "move_file":       lambda src, dst: __import__('app.services.file_ops', fromlist=['move_file']).move_file(src, dst),
    "delete_file":     lambda path: __import__('app.services.file_ops', fromlist=['delete_file']).delete_file(path),
    "search_files":     lambda name, root_dir='home': __import__('app.services.file_ops', fromlist=['search_files']).search_files(name, root_dir),
    "create_folder":    lambda path: __import__('app.services.file_ops', fromlist=['create_folder']).create_folder(path),
    "bulk_rename":     lambda directory, find, replace: __import__('app.services.file_ops', fromlist=['bulk_rename']).bulk_rename(directory, find, replace),
    "diff_files":       lambda path1, path2: __import__('app.services.file_ops', fromlist=['diff_files']).diff_files(path1, path2),
    # Gmail integration (Step 5 - Browser based)
    "check_emails":     lambda query='is:unread', max_results=5: __import__('app.services.browser_mail', fromlist=['check_emails']).check_emails(query, max_results),
    "list_unread":     lambda max_results=5: __import__('app.services.browser_mail', fromlist=['list_unread']).list_unread(max_results),
    "get_email_body":   lambda email_id: __import__('app.services.browser_mail', fromlist=['get_email_body']).get_email_body(email_id),
    "summarize_inbox":  lambda max_results=10: __import__('app.services.browser_mail', fromlist=['summarize_inbox']).summarize_inbox(max_results),
    "smart_mail_action": lambda task: __import__('app.services.browser_mail', fromlist=['smart_mail_action']).smart_mail_action(task),
    # Browser automation (Step 6)
    "browse_and_read":  lambda url: __import__('app.services.browser_tool', fromlist=['browse_and_read']).browse_and_read(url),
    "search_on_site":   lambda site_url, query: __import__('app.services.browser_tool', fromlist=['search_on_site']).search_on_site(site_url, query),
    "click_element":    lambda page_url, text: __import__('app.services.browser_tool', fromlist=['click_element']).click_element(page_url, text),
    "scroll_and_read":  lambda url, px=1000: __import__('app.services.browser_tool', fromlist=['scroll_and_read']).scroll_and_read(url, px),
    # Google Calendar (Step 8)
    "get_upcoming_events": lambda days=7: __import__('app.services.calendar_tool', fromlist=['get_upcoming_events']).get_upcoming_events(days),
    "check_today_schedule": lambda: __import__('app.services.calendar_tool', fromlist=['check_today_schedule']).check_today_schedule(),
    "add_event":       lambda title, date, time=None, notes="": __import__('app.services.calendar_tool', fromlist=['add_event']).add_event(title, date, time, notes),
    # Persistent Memory (Step 10)
    "save_fact":       lambda topic, fact: __import__('app.services.memory_tool', fromlist=['save_fact']).save_fact(topic, fact),
    "recall_facts":     lambda topic=None: __import__('app.services.memory_tool', fromlist=['recall_facts']).recall_facts(topic),
    "get_morning_brief": lambda: __import__('app.services.memory_tool', fromlist=['get_morning_brief']).get_morning_brief(),
    "update_fact":     lambda topic, old_fact, new_fact: __import__('app.services.memory_tool', fromlist=['update_fact']).update_fact(topic, old_fact, new_fact),
    "forget_fact":     lambda topic: __import__('app.services.memory_tool', fromlist=['forget_fact']).forget_fact(topic),
    # Extended browser tools (Step 6 enhanced)
    "fill_form":       lambda url, fields: __import__('app.services.browser_tool', fromlist=['fill_form']).fill_form(url, fields),
    "browse_and_paginate": lambda url, pages=3: __import__('app.services.browser_tool', fromlist=['browse_and_paginate']).browse_and_paginate(url, pages),
    "smart_web_action": lambda site_name, task: __import__('app.services.agentic_web', fromlist=['agentic_web_action']).agentic_web_action(site_name, task),
    "agentic_web_action": lambda site_or_task, specific_task=None: __import__('app.services.agentic_web', fromlist=['agentic_web_action']).agentic_web_action(site_or_task, specific_task),
    # ── UIA-powered Windows UI Automation tools ──────────────────────────────
    # These operate WITHOUT moving the mouse cursor.
    # They use the same API as Windows Narrator / screen readers.
    "click_ui_element_uia":  click_ui_element_uia,
    "type_into_ui_element":  type_into_ui_element,
    "read_ui_element_text":  read_ui_element_text,
    "dump_app_ui_tree":      dump_app_ui_tree,
    # ── Assignment Automation Tools ──────────────────────────────────────────
    # Isolated in app/services/assignment_tool.py + assignment_answers.py
    # Accessible via both frontend UI and voice through unified /chat route.
    # Phase 1: Extraction
    "extract_questions":   lambda pdf_path: __import__('app.services.assignment_tool', fromlist=['extract_questions']).extract_questions(pdf_path),
    "list_assignments":    lambda: __import__('app.services.assignment_tool', fromlist=['list_assignments']).list_assignments(),
    # Phase 2: Answer Generation (browser: Gemini→ChatGPT→DeepSeek, fallback: Groq API)
    "generate_answers":    lambda questions_json, pdf_path='': __import__('app.services.assignment_answers', fromlist=['generate_answers']).generate_answers(questions_json, pdf_path),
    "generate_answer":     lambda question, question_type='long_answer', has_figure=False: __import__('app.services.assignment_answers', fromlist=['generate_answer']).generate_answer(question, question_type, has_figure),
    # Phase 3: Answer Humanization (browser: Paraphraser/Scribbr/QuillBot, fallback: Groq API)
    "humanize_all_answers": lambda qa_json: __import__('app.services.assignment_humanizer', fromlist=['humanize_all_answers']).humanize_all_answers(qa_json),
    "humanize_text":       lambda text, force_site='': __import__('app.services.assignment_humanizer', fromlist=['humanize_text']).humanize_text(text, force_site),
    # Phase 4: Document Assembly (Word / PPTX)
    "assemble_assignment": lambda qa_json, filename, format_type='word': __import__('app.services.assignment_assembler', fromlist=['assemble_assignment']).assemble_assignment(qa_json, filename, format_type),
    # Phase 5: Master Orchestrator Pipeline
    "do_assignment":       lambda pdf_path, output_format='word', humanize=True: __import__('app.services.assignment_pipeline', fromlist=['do_assignment']).do_assignment(pdf_path, output_format, humanize),
    # ── Syllabus Auditor (Stateless Tool) ────────────────────────────────────
    "audit_playlist_syllabus": lambda playlist_url, image_path: __import__('app.services.syllabus_auditor', fromlist=['audit_playlist_syllabus']).audit_playlist_syllabus(playlist_url, image_path),
    # ── AI Content Humanizer (5-Stage Pipeline) ──────────────────────────────
    "humanize_ai_content": lambda text: __import__('app.services.content_humanizer', fromlist=['humanize_text_sync']).humanize_text_sync(text),
    # ── PowerPoint AI (ppt_tool.py + research_pipeline.py) ───────────────────
    
    # Phase 1: Pure PPT Generation
    "ppt_create": _ppt_create,
    "ppt_edit":   _ppt_edit,
    "ppt_styles": _ppt_styles,
    
    # Phase 3: Research-backed PPT Generation
    "research_and_create_ppt": _research_and_create_ppt,
    # ── Prompt Enhancer (Step 1) ─────────────────────────────────────────────
    "enhance_prompt": lambda raw_prompt: __import__('app.services.skill_prompt_enhancer', fromlist=['enhance_prompt']).enhance_prompt(raw_prompt),
    # ── DSA Mode (Leetcode Enforcer) ─────────────────────────────────────────
    "activate_dsa_mode": lambda num_questions: __import__('app.services.dsa_enforcer', fromlist=['get_dsa_enforcer']).get_dsa_enforcer().start_mode(int(num_questions)),
    "deactivate_dsa_mode": lambda: __import__('app.services.dsa_enforcer', fromlist=['get_dsa_enforcer']).get_dsa_enforcer().stop_mode(),
    "dsa_status": get_dsa_cache_status,
    # ── Media Enhancement Tool ───────────────────────────────────────────────
    "enhance_media": lambda file_path: __import__('app.services.media_enhancement', fromlist=['enhance_media']).enhance_media(file_path),
    # ── Social Content Manager ───────────────────────────────────────────────
    "generate_social_content": lambda idea, platform="Instagram", tone="engaging", creativity=50.0, formality=50.0, smart_emojis=True, auto_hashtag=True, contextual_suggestions=True, target_audience="": __import__('app.services.social_content_manager', fromlist=['generate_social_content']).generate_social_content(idea, platform, tone, creativity, formality, smart_emojis, auto_hashtag, contextual_suggestions, target_audience),
    "refine_social_content": lambda original_content, refinement_instruction, platform="Instagram": __import__('app.services.social_content_manager', fromlist=['refine_social_content']).refine_social_content(original_content, refinement_instruction, platform),
    # ── Long-Term RAG Memory (MySQL + FAISS) ────────────────────────────────
    # Isolated per Rule #1. Accessible via both frontend and voice through /chat.
    # The LLM router can call this when user asks explicit memory questions.
    "recall_memory": lambda query: __import__('app.services.rag_memory', fromlist=['recall', 'format_recall_for_prompt']).__dict__,  # handled async in chat.py
    # ── Neural Cache Tools ───────────────────────────────────────────────────
    # Rule #1: isolated in neural_cache/ package, no cross-tool imports.
    # Rule #4: accessible via both frontend and voice through unified /chat.
    "cache_set": cache_set,
    "cache_get": cache_get,
    
    # ── Air Drawing Tool ─────────────────────────────────────────────────────
    "open_air_drawing": lambda: __import__('app.services.air_drawing_tool', fromlist=['open_air_drawing']).open_air_drawing(),
}